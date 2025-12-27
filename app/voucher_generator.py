"""Voucher generator module.

This module generates voucher documents from parsed ORGA data by filling
the blank voucher template while preserving the original formatting.
"""
import logging
import os
import copy
import tempfile
from datetime import date, timedelta
from typing import List, Tuple, Any, Optional
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .models import (
    ParsedORGA, HotelStay, TransferVoucher, TransferLeg,
    ActivityVoucher, RestaurantVoucher, CarRentalVoucher, GolfVoucher
)
from .supplier_info import get_supplier_info

logger = logging.getLogger(__name__)

# Colors from the template
COLOR_GRAY = RGBColor(0x74, 0x74, 0x74)
COLOR_RED = RGBColor(0xEE, 0x00, 0x00)
COLOR_DARK = RGBColor(0x22, 0x22, 0x22)


def format_date(d: date) -> str:
    """Format date as 'DD Month YYYY'."""
    return d.strftime("%d %B %Y")


def format_date_short(d: date) -> str:
    """Format date as 'DD.MM.YYYY'."""
    return d.strftime("%d.%m.%Y")


def get_board_basis_text(board: str) -> str:
    """Convert board abbreviation to full text."""
    board_map = {
        "RO": "Room Only",
        "BB": "Bed & Breakfast",
        "HB": "Half Board",
        "FB": "Full Board",
        "FB+": "Full Board Plus - Dinner, Bed, Breakfast, Lunch and Activities",
        "AI": "All Inclusive"
    }
    return board_map.get(board.upper().strip(), board)


def find_paragraph_by_start(cell, start_text: str):
    """Find a paragraph that starts with specific text."""
    for para in cell.paragraphs:
        if para.text.strip().upper().startswith(start_text.upper()):
            return para
    return None


def add_text_with_style(paragraph, text: str, bold=False, italic=False, color=None, size=None):
    """Add a run of text with specific styling."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return run


def clear_paragraph_after_label(paragraph, label: str):
    """Clear text after a label while preserving the label formatting."""
    # Keep the label, remove everything after
    full_text = paragraph.text
    if label in full_text:
        # We need to rebuild the paragraph
        pass


class VoucherGenerator:
    """Generates voucher documents from parsed ORGA data."""
    
    def __init__(self, template_path: str):
        """Initialize with path to blank voucher template."""
        self.template_path = template_path
        self.generated_vouchers: List[Tuple[str, str, date]] = []
    
    def generate_all(
        self,
        parsed_data: ParsedORGA,
        traveller_names: str,
        ref_no: str = "",
        group_text: str = "",
        output_dir: str = None
    ) -> List[Tuple[str, str, date]]:
        """Generate all vouchers from parsed ORGA data."""
        if output_dir is None:
            output_dir = tempfile.mkdtemp(prefix="vouchers_")
        
        os.makedirs(output_dir, exist_ok=True)
        self.generated_vouchers = []
        
        # Generate hotel vouchers
        for i, hotel in enumerate(parsed_data.hotels):
            path = os.path.join(output_dir, f"hotel_{i+1}_{self._safe_filename(hotel.supplier)}.docx")
            self._generate_hotel_voucher(hotel, traveller_names, ref_no, path)
            self.generated_vouchers.append((path, "hotel", hotel.check_in))
            logger.info(f"Generated hotel voucher: {hotel.supplier}")
        
        # Generate transfer vouchers
        for i, transfer in enumerate(parsed_data.transfers):
            path = os.path.join(output_dir, f"transfer_{i+1}_{self._safe_filename(transfer.supplier)}.docx")
            self._generate_transfer_voucher(transfer, traveller_names, ref_no, path)
            earliest = min(leg.date for leg in transfer.legs) if transfer.legs else date.today()
            self.generated_vouchers.append((path, "transfer", earliest))
            logger.info(f"Generated transfer voucher: {transfer.supplier}")
        
        # Generate car rental vouchers
        for i, car in enumerate(parsed_data.car_rentals):
            path = os.path.join(output_dir, f"car_rental_{i+1}_{self._safe_filename(car.supplier)}.docx")
            self._generate_car_rental_voucher(car, traveller_names, ref_no, group_text, path)
            self.generated_vouchers.append((path, "car_rental", car.pickup_date))
            logger.info(f"Generated car rental voucher: {car.supplier}")
        
        # Generate activity vouchers
        for i, activity in enumerate(parsed_data.activities):
            path = os.path.join(output_dir, f"activity_{i+1}_{self._safe_filename(activity.supplier)}.docx")
            self._generate_activity_voucher(activity, traveller_names, ref_no, path)
            earliest = min(e.date for e in activity.entries) if activity.entries else date.today()
            self.generated_vouchers.append((path, "activity", earliest))
            logger.info(f"Generated activity voucher: {activity.supplier}")
        
        # Generate restaurant vouchers
        for i, restaurant in enumerate(parsed_data.restaurants):
            path = os.path.join(output_dir, f"restaurant_{i+1}_{self._safe_filename(restaurant.supplier)}.docx")
            self._generate_restaurant_voucher(restaurant, traveller_names, ref_no, path)
            self.generated_vouchers.append((path, "restaurant", restaurant.date))
            logger.info(f"Generated restaurant voucher: {restaurant.supplier}")
        
        # Generate golf vouchers
        for i, golf in enumerate(parsed_data.golf):
            path = os.path.join(output_dir, f"golf_{i+1}_{self._safe_filename(golf.supplier)}.docx")
            self._generate_golf_voucher(golf, traveller_names, ref_no, path)
            self.generated_vouchers.append((path, "golf", golf.date))
            logger.info(f"Generated golf voucher: {golf.supplier}")
        
        return self.generated_vouchers
    
    def _safe_filename(self, name: str) -> str:
        """Create a safe filename from supplier name."""
        safe = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_'))
        return safe.strip().replace(' ', '_')[:50]
    
    def _load_template(self) -> Document:
        """Load a fresh copy of the template document."""
        return Document(self.template_path)
    
    def _get_supplier_info(self, supplier_name: str) -> dict:
        """Get supplier info from database."""
        return get_supplier_info(supplier_name)
    
    def _fill_supplier_header(self, doc: Document, supplier_name: str):
        """Fill the supplier header section (Row 1 of the table)."""
        if not doc.tables:
            return
        
        table = doc.tables[0]
        if len(table.rows) < 2:
            return
        
        cell = table.rows[1].cells[0]
        info = self._get_supplier_info(supplier_name)
        
        # Clear existing content
        for para in cell.paragraphs:
            para.clear()
        
        # Add supplier name (bold, larger)
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = p.add_run(info.get("display_name", supplier_name.upper()))
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_DARK
        
        # Add address
        if info.get("address"):
            p = cell.add_paragraph()
            run = p.add_run(info["address"])
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_GRAY
        
        # Add phone
        if info.get("phone"):
            p = cell.add_paragraph()
            run = p.add_run(f"Tel: {info['phone']}")
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_GRAY
        
        # Add GPS
        if info.get("gps"):
            p = cell.add_paragraph()
            run = p.add_run(f"GPS: {info['gps']}")
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_GRAY
    
    def _fill_content_section(
        self,
        doc: Document,
        traveller_names: str,
        ref_no: str = "",
        group_text: str = "",
        check_in: str = "",
        check_out: str = "",
        nights: int = 0,
        date_single: str = "",
        time_text: str = "",
        included_services: List[str] = None,
        notes: str = ""
    ):
        """Fill the main content section (Row 2 of the table)."""
        if not doc.tables:
            return
        
        table = doc.tables[0]
        if len(table.rows) < 3:
            return
        
        cell = table.rows[2].cells[0]
        
        # Clear existing paragraphs
        for para in cell.paragraphs:
            para.clear()
        
        paragraphs = cell.paragraphs
        current_idx = 0
        
        def get_or_add_para():
            nonlocal current_idx
            if current_idx < len(paragraphs):
                p = paragraphs[current_idx]
            else:
                p = cell.add_paragraph()
            current_idx += 1
            return p
        
        # Empty line at start
        get_or_add_para()
        
        # TRAVELLERS line
        p = get_or_add_para()
        run = p.add_run("TRAVELLERS: ")
        run.bold = True
        run.font.color.rgb = COLOR_GRAY
        run = p.add_run(traveller_names)
        run.font.color.rgb = COLOR_GRAY
        
        # Empty line
        get_or_add_para()
        
        # REF NO line
        p = get_or_add_para()
        run = p.add_run("REF NO: ")
        run.bold = True
        run.font.color.rgb = COLOR_GRAY
        run = p.add_run(ref_no if ref_no else "")
        run.font.color.rgb = COLOR_GRAY
        
        # Empty line
        get_or_add_para()
        
        # GROUP line (if provided)
        if group_text:
            p = get_or_add_para()
            run = p.add_run("GROUP: ")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run(group_text)
            run.font.color.rgb = COLOR_GRAY
            get_or_add_para()
        
        # CHECK IN / CHECK OUT lines (for hotels)
        if check_in and check_out:
            p = get_or_add_para()
            run = p.add_run("CHECK IN: ")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run(f"{check_in}")
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run("                TIME: ")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run("14h00")
            run.font.color.rgb = COLOR_GRAY
            
            p = get_or_add_para()
            run = p.add_run("CHECK OUT: ")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run(f"{check_out}")
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run("              TIME: ")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run("11h00")
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run("              NIGHTS: ")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run(str(nights))
            run.font.color.rgb = COLOR_GRAY
            
            get_or_add_para()
        
        # DATE line (for activities/restaurants)
        elif date_single:
            p = get_or_add_para()
            run = p.add_run("DATE: ")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run(date_single)
            run.font.color.rgb = COLOR_GRAY
            
            get_or_add_para()
        
        # TIME line (for activities/restaurants)
        if time_text:
            p = get_or_add_para()
            run = p.add_run("TIME: ")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            run = p.add_run(time_text)
            run.font.color.rgb = COLOR_GRAY
            
            get_or_add_para()
        
        # Included Services section
        p = get_or_add_para()
        run = p.add_run("Included Services:")
        run.bold = True
        run.italic = True
        run.font.color.rgb = COLOR_GRAY
        
        get_or_add_para()
        
        # Service items as bullet points
        if included_services:
            for service in included_services:
                if service.strip():
                    p = get_or_add_para()
                    # Add bullet point
                    run = p.add_run("•    ")
                    run.font.color.rgb = COLOR_GRAY
                    
                    # Check if service has a label (e.g., "Accommodation Type:")
                    if ":" in service:
                        parts = service.split(":", 1)
                        run = p.add_run(f"{parts[0]}:")
                        run.bold = True
                        run.italic = True
                        run.font.color.rgb = COLOR_GRAY
                        if len(parts) > 1:
                            run = p.add_run(f" {parts[1].strip()}")
                            run.font.color.rgb = COLOR_GRAY
                    else:
                        run = p.add_run(service)
                        run.font.color.rgb = COLOR_GRAY
        
        # Notes section
        if notes:
            get_or_add_para()
            p = get_or_add_para()
            run = p.add_run("Notes:")
            run.bold = True
            run.font.color.rgb = COLOR_GRAY
            
            p = get_or_add_para()
            run = p.add_run(notes)
            run.font.color.rgb = COLOR_GRAY
        
        # Empty line before disclaimer
        get_or_add_para()
        
        # Red disclaimer line
        p = get_or_add_para()
        run = p.add_run("All additional services are for guest's own account")
        run.bold = True
        run.italic = True
        run.font.color.rgb = COLOR_RED
        
        # Final empty line
        get_or_add_para()
    
    def _generate_hotel_voucher(
        self,
        hotel: HotelStay,
        traveller_names: str,
        ref_no: str,
        output_path: str
    ):
        """Generate a hotel voucher."""
        doc = self._load_template()
        
        # Fill supplier header
        self._fill_supplier_header(doc, hotel.supplier)
        
        # Build included services list
        services = []
        room_text = f"Accommodation Type: X1 {hotel.room_type} - DBL" if hotel.room_type else "Accommodation Type: Double Room"
        services.append(room_text)
        
        if hotel.board:
            services.append(f"Board Basis: {get_board_basis_text(hotel.board)}")
        
        # For safari lodges (FB+ or FB), add activities
        if hotel.board and hotel.board.upper() in ["FB+", "FB"]:
            services.append("")
            services.append("Activities:")
            current = hotel.check_in
            while current < hotel.check_out:
                if current == hotel.check_in:
                    services.append(f"    {format_date_short(current)} – X1 Afternoon Game Drive")
                elif current == hotel.check_out - timedelta(days=1):
                    services.append(f"    {format_date_short(current)} – X1 Morning Game Drive")
                else:
                    services.append(f"    {format_date_short(current)} – X1 Morning & Afternoon Game Drive")
                current += timedelta(days=1)
        
        # Fill content
        self._fill_content_section(
            doc=doc,
            traveller_names=traveller_names,
            ref_no=ref_no,
            check_in=format_date(hotel.check_in),
            check_out=format_date(hotel.check_out),
            nights=hotel.nights,
            included_services=services,
            notes=hotel.notes
        )
        
        doc.save(output_path)
    
    def _generate_transfer_voucher(
        self,
        transfer: TransferVoucher,
        traveller_names: str,
        ref_no: str,
        output_path: str
    ):
        """Generate a transfer voucher."""
        doc = self._load_template()
        
        self._fill_supplier_header(doc, transfer.supplier)
        
        # Build services list
        services = []
        for leg in transfer.legs:
            pickup_text = f"Pick Up: {format_date_short(leg.date)}"
            if leg.pickup_location:
                pickup_text += f" – {leg.pickup_location}"
            if leg.pickup_time:
                pickup_text += f" @ {leg.pickup_time}"
            if leg.flight_number:
                pickup_text += f" (Flight {leg.flight_number})"
            if "airport" in (leg.pickup_location or "").lower():
                pickup_text += " – Your driver will meet you in the arrivals hall with your name board."
            
            services.append(pickup_text)
            
            if leg.dropoff_location:
                services.append(f"Drop Off: {leg.dropoff_location}")
            
            services.append("")  # Blank line between legs
        
        # Collect notes
        notes = transfer.notes
        for leg in transfer.legs:
            if leg.notes:
                notes = (notes + "\n" + leg.notes).strip() if notes else leg.notes
        
        self._fill_content_section(
            doc=doc,
            traveller_names=traveller_names,
            ref_no=ref_no,
            included_services=services,
            notes=notes
        )
        
        doc.save(output_path)
    
    def _generate_car_rental_voucher(
        self,
        car: CarRentalVoucher,
        traveller_names: str,
        ref_no: str,
        group_text: str,
        output_path: str
    ):
        """Generate a car rental voucher."""
        doc = self._load_template()
        
        self._fill_supplier_header(doc, car.supplier)
        
        # Car group info as the GROUP field
        car_group_lines = [
            car.car_group,
            "Unlimited Mileage",
            "Zero Excess",
            "Including glass and tire insurance",
            "Full to Full Fuel Policy"
        ]
        group_info = "\n".join(car_group_lines)
        
        services = [
            f"Pick Up: {format_date_short(car.pickup_date)} – {car.pickup_location}",
            f"Drop Off: {format_date_short(car.dropoff_date)} – {car.dropoff_location}"
        ]
        
        self._fill_content_section(
            doc=doc,
            traveller_names=traveller_names,
            ref_no=ref_no,
            group_text=group_info,
            included_services=services,
            notes=car.notes
        )
        
        doc.save(output_path)
    
    def _generate_activity_voucher(
        self,
        activity: ActivityVoucher,
        traveller_names: str,
        ref_no: str,
        output_path: str
    ):
        """Generate an activity voucher."""
        doc = self._load_template()
        
        self._fill_supplier_header(doc, activity.supplier)
        
        if len(activity.entries) == 1:
            entry = activity.entries[0]
            services = [entry.activity_name]
            date_single = format_date(entry.date)
            time_text = entry.time if entry.time else ""
            notes = entry.notes
        else:
            services = []
            for entry in activity.entries:
                line = f"{format_date_short(entry.date)}"
                if entry.time:
                    line += f" – {entry.time}"
                line += f" – {entry.activity_name}"
                services.append(line)
            date_single = ""
            time_text = ""
            notes = "\n".join(e.notes for e in activity.entries if e.notes)
        
        self._fill_content_section(
            doc=doc,
            traveller_names=traveller_names,
            ref_no=ref_no,
            date_single=date_single,
            time_text=time_text,
            included_services=services,
            notes=notes
        )
        
        doc.save(output_path)
    
    def _generate_restaurant_voucher(
        self,
        restaurant: RestaurantVoucher,
        traveller_names: str,
        ref_no: str,
        output_path: str
    ):
        """Generate a restaurant voucher."""
        doc = self._load_template()
        
        self._fill_supplier_header(doc, restaurant.supplier)
        
        services = [restaurant.notes if restaurant.notes else "Dinner reservation"]
        
        self._fill_content_section(
            doc=doc,
            traveller_names=traveller_names,
            ref_no=ref_no,
            date_single=format_date(restaurant.date),
            time_text=restaurant.time if restaurant.time else "",
            included_services=services
        )
        
        doc.save(output_path)
    
    def _generate_golf_voucher(
        self,
        golf: GolfVoucher,
        traveller_names: str,
        ref_no: str,
        output_path: str
    ):
        """Generate a golf voucher."""
        doc = self._load_template()
        
        self._fill_supplier_header(doc, golf.supplier)
        
        services = [f"Golf Course: {golf.course}"]
        if golf.cart:
            services.append(f"Cart: {golf.cart}")
        if golf.rental_set:
            services.append(f"Rental Set: {golf.rental_set}")
        
        self._fill_content_section(
            doc=doc,
            traveller_names=traveller_names,
            ref_no=ref_no,
            date_single=format_date(golf.date),
            time_text=f"Tee Time: {golf.tee_time}" if golf.tee_time else "",
            included_services=services,
            notes=golf.notes
        )
        
        doc.save(output_path)
